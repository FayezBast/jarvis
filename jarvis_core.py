import os
import subprocess
import json
import uuid
import webbrowser
import requests
import pyautogui
import pyperclip
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from ai_core import AI_Core, CommandAnalysis
from config import Config
from security import SecurityValidator
from logger import log_info, log_error, log_warning

class JarvisCore:
    """Redesigned JARVIS Core with comprehensive action execution"""
    
    def __init__(self, voice_io):
        log_info("Initializing JARVIS Core...")
        
        # Core components
        api_key = os.getenv('GOOGLE_API_KEY')
        self.ai_core = AI_Core(api_key)
        self.voice_io = voice_io
        
        # Session management
        self.conversation_history = []
        self.session_id = str(uuid.uuid4())
        
        # Memory system
        self.memory_file = "jarvis_memory.json"
        self.memory = self._load_memory()
        
        # Ensure workspace directory exists
        Config.WORKSPACE_DIR.mkdir(exist_ok=True)
        
        # Initialize pyautogui with safety
        pyautogui.FAILSAFE = True
        pyautogui.PAUSE = 0.1
        
        log_info("JARVIS Core initialized successfully.")

    def process_command(self, command: str) -> str:
        """Main command processing with comprehensive execution"""
        log_info(f"Processing command: '{command}'")
        
        if not command.strip():
            response = "How can I help you?"
            self.voice_io.speak(response)
            return response

        try:
            # Analyze command
            analysis = self.ai_core.analyze_command(command, self.conversation_history)
            log_info(f"Analysis: intent={analysis.intent}, action={analysis.action}, params={analysis.parameters}")
            
            # Execute based on analysis
            if analysis.response:
                # Direct response from AI (conversation mode)
                response = analysis.response
            else:
                # Execute the action
                response = self._execute_action(analysis)
            
            # Update conversation history
            self._update_conversation_history(command, response)
            
            # Store in memory
            self._store_interaction(command, response, analysis)
            
            # Speak the response
            self.voice_io.speak(response)
            return response

        except Exception as e:
            log_error(f"Command processing failed: {e}", exc_info=True)
            error_message = "I encountered an error while processing your request."
            self.voice_io.speak(error_message)
            return error_message

    def _execute_action(self, analysis: CommandAnalysis) -> str:
        """Execute action based on analysis with comprehensive coverage"""
        
        intent = analysis.intent
        action = analysis.action
        params = analysis.parameters
        
        # Route to appropriate handler
        try:
            if intent == 'file_creation':
                return self._handle_file_creation(action, params)
            elif intent == 'file_management':
                return self._handle_file_management(action, params)
            elif intent == 'system_control':
                return self._handle_system_control(action, params)
            elif intent == 'web_browse':
                return self._handle_web_browse(action, params)
            elif intent == 'powershell_task':
                return self._handle_powershell_task(action, params)
            elif intent == 'weather_inquiry':
                return self._handle_weather_inquiry(action, params)
            elif intent == 'knowledge_inquiry':
                return self._handle_knowledge_inquiry(action, params)
            elif intent == 'memory_query':
                return self._handle_memory_query(action, params)
            elif intent == 'clipboard_management':
                return self._handle_clipboard_management(action, params)
            elif intent == 'windows_automation':
                return self._handle_windows_automation(action, params)
            elif intent == 'system_status':
                return self._handle_system_status(action, params)
            elif intent == 'network_operations':
                return self._handle_network_operations(action, params)
            elif intent == 'media_control':
                return self._handle_media_control(action, params)
            elif intent == 'help':
                return self._get_comprehensive_help()
            else:
                return f"I understand you want to {intent}, but I'm not sure how to handle '{action}' yet."
                
        except Exception as e:
            log_error(f"Action execution failed for {intent}.{action}: {e}")
            return f"I encountered an error while trying to {action}. Please try rephrasing your request."

    # === FILE CREATION HANDLERS ===
    def _handle_file_creation(self, action: str, params: Dict) -> str:
        """Handle file creation with various formats"""
        try:
            file_type = params.get('file_type', 'txt')
            topic = params.get('content_topic', params.get('topic', 'general'))
            filename = params.get('filename', f"{topic.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{file_type}")
            
            # Sanitize filename
            filename = SecurityValidator.validate_filename(filename)
            file_path = Config.WORKSPACE_DIR / filename
            
            # Generate content
            content = self.ai_core.generate_file_content(topic, file_type)
            
            # Create file based on type
            if file_type == 'docx':
                return self._create_word_document(file_path, content, topic)
            elif file_type == 'xlsx':
                return self._create_excel_document(file_path, content, topic)
            elif file_type == 'pdf':
                return self._create_pdf_document(file_path, content, topic)
            else:
                # Text-based files (txt, py, json, etc.)
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return f"Created {file_type.upper()} file: {filename}"
                
        except Exception as e:
            log_error(f"File creation failed: {e}")
            return f"Failed to create file: {str(e)}"

    def _create_word_document(self, file_path: Path, content: str, topic: str) -> str:
        """Create Word document using python-docx if available"""
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(topic.title(), 0)
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('#'):
                        # Handle markdown headers
                        level = para.count('#')
                        text = para.replace('#', '').strip()
                        doc.add_heading(text, level)
                    else:
                        doc.add_paragraph(para.strip())
            
            doc.save(file_path)
            return f"Created Word document: {file_path.name}"
            
        except ImportError:
            log_warning("python-docx not available, creating as text file")
            with open(file_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                f.write(f"{topic.title()}\n{'='*len(topic)}\n\n{content}")
            return f"Created text file instead: {file_path.with_suffix('.txt').name}"

    def _create_excel_document(self, file_path: Path, content: str, topic: str) -> str:
        """Create Excel document using openpyxl if available"""
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = topic.title()[:31]  # Excel sheet name limit
            
            ws['A1'] = topic.title()
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            
            # Add basic content
            lines = content.split('\n')
            for i, line in enumerate(lines[:50], 2):  # Limit to 50 rows
                if line.strip():
                    ws[f'A{i}'] = line.strip()
            
            wb.save(file_path)
            return f"Created Excel document: {file_path.name}"
            
        except ImportError:
            log_warning("openpyxl not available, creating as CSV file")
            with open(file_path.with_suffix('.csv'), 'w', encoding='utf-8') as f:
                f.write(f"{topic.title()}\n{content}")
            return f"Created CSV file instead: {file_path.with_suffix('.csv').name}"

    def _create_pdf_document(self, file_path: Path, content: str, topic: str) -> str:
        """Create PDF document using reportlab if available"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            c = canvas.Canvas(str(file_path), pagesize=letter)
            width, height = letter
            
            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, topic.title())
            
            # Content
            c.setFont("Helvetica", 12)
            y_position = height - 100
            
            lines = content.split('\n')
            for line in lines:
                if y_position < 50:  # New page
                    c.showPage()
                    y_position = height - 50
                
                if line.strip():
                    c.drawString(50, y_position, line.strip()[:80])  # Limit line length
                    y_position -= 15
            
            c.save()
            return f"Created PDF document: {file_path.name}"
            
        except ImportError:
            log_warning("reportlab not available, creating as text file")
            with open(file_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                f.write(f"{topic.title()}\n{'='*len(topic)}\n\n{content}")
            return f"Created text file instead: {file_path.with_suffix('.txt').name}"

    # === FILE MANAGEMENT HANDLERS ===
    def _handle_file_management(self, action: str, params: Dict) -> str:
        """Handle file management operations"""
        try:
            if action == 'list_files':
                directory = params.get('directory', str(Config.WORKSPACE_DIR))
                return self._list_files(directory)
            elif action == 'read_file':
                file_path = params.get('file_path', '')
                return self._read_file(file_path)
            elif action == 'delete_file':
                file_path = params.get('file_path', '')
                if not file_path:
                    # Try to extract from command
                    return "Please specify the complete file path to delete."
                return self._delete_file(file_path)
            elif action == 'copy_file':
                source = params.get('source', '')
                destination = params.get('destination', '')
                return self._copy_file(source, destination)
            elif action == 'open_file':
                file_path = params.get('file_path', '')
                if not file_path:
                    # Default to opening workspace folder
                    file_path = str(Config.WORKSPACE_DIR)
                return self._open_file(file_path)
            elif action == 'execute_file':
                file_path = params.get('file_path', '')
                return self._execute_file(file_path)
            else:
                return f"File management action '{action}' not implemented yet."
        except Exception as e:
            return f"File operation failed: {str(e)}"

    def _list_files(self, directory: str) -> str:
        """List files in directory"""
        try:
            dir_path = Path(directory)
            if not dir_path.exists():
                return f"Directory not found: {directory}"
            
            files = list(dir_path.iterdir())
            if not files:
                return f"No files found in {directory}"
            
            file_list = []
            for file in files[:20]:  # Limit to 20 files
                if file.is_file():
                    size = file.stat().st_size
                    modified = datetime.fromtimestamp(file.stat().st_mtime).strftime('%Y-%m-%d %H:%M')
                    file_list.append(f"{file.name} ({size} bytes, modified: {modified})")
                elif file.is_dir():
                    file_list.append(f"{file.name}/ (directory)")
            
            return f"Files in {directory}:\n" + "\n".join(file_list)
        except Exception as e:
            return f"Failed to list files: {str(e)}"

    def _read_file(self, file_path: str) -> str:
        """Read file content"""
        try:
            if not file_path:
                return "Please specify a file path."
            
            path = Path(file_path)
            if not path.is_absolute():
                path = Config.WORKSPACE_DIR / path
            
            if not path.exists():
                return f"File not found: {file_path}"
            
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read(2000)  # Limit to 2000 characters
            
            if len(content) >= 2000:
                content += "... (truncated)"
            
            return f"Content of {path.name}:\n{content}"
        except Exception as e:
            return f"Failed to read file: {str(e)}"

    def _delete_file(self, file_path: str) -> str:
        """Delete a file"""
        try:
            if not file_path:
                return "Please specify a file path."
            
            path = Path(file_path)
            if not path.is_absolute():
                path = Config.WORKSPACE_DIR / path
            
            if not path.exists():
                return f"File not found: {file_path}"
            
            path.unlink()
            return f"File deleted: {path.name}"
        except Exception as e:
            return f"Failed to delete file: {str(e)}"

    def _copy_file(self, source: str, destination: str) -> str:
        """Copy a file"""
        try:
            import shutil
            
            src_path = Path(source)
            if not src_path.is_absolute():
                src_path = Config.WORKSPACE_DIR / src_path
            
            dst_path = Path(destination)
            if not dst_path.is_absolute():
                dst_path = Config.WORKSPACE_DIR / dst_path
            
            shutil.copy2(src_path, dst_path)
            return f"File copied from {src_path.name} to {dst_path.name}"
        except Exception as e:
            return f"Failed to copy file: {str(e)}"

    def _open_file(self, file_path: str) -> str:
        """Open a file with default application"""
        try:
            path = Path(file_path)
            if not path.is_absolute():
                path = Config.WORKSPACE_DIR / path
            
            if not path.exists():
                return f"File not found: {file_path}"
            
            os.startfile(str(path))
            return f"Opening file: {path.name}"
        except Exception as e:
            return f"Failed to open file: {str(e)}"

    def _execute_file(self, file_path: str) -> str:
        """Execute a file (Python scripts, etc.)"""
        try:
            if not file_path:
                return "Please specify a file path to execute."
            
            path = Path(file_path)
            if not path.is_absolute():
                path = Config.WORKSPACE_DIR / path
            
            if not path.exists():
                return f"File not found: {file_path}"
            
            if path.suffix.lower() == '.py':
                # Execute Python script
                result = subprocess.run([
                    'python', str(path)
                ], capture_output=True, text=True, timeout=30)
                
                output = result.stdout.strip()
                if result.stderr:
                    output += f"\nErrors: {result.stderr.strip()}"
                
                return f"Executed {path.name}:\n{output}" if output else f"Executed {path.name} successfully."
            
            elif path.suffix.lower() in ['.bat', '.cmd']:
                # Execute batch file
                result = subprocess.run([str(path)], capture_output=True, text=True, timeout=30, shell=True)
                output = result.stdout.strip()
                return f"Executed {path.name}:\n{output}" if output else f"Executed {path.name} successfully."
            
            else:
                # Try to open with default application
                os.startfile(str(path))
                return f"Opened {path.name} with default application."
                
        except subprocess.TimeoutExpired:
            return f"Execution timed out after 30 seconds."
        except Exception as e:
            return f"Failed to execute file: {str(e)}"

    # === SYSTEM CONTROL HANDLERS ===
    def _handle_system_control(self, action: str, params: Dict) -> str:
        """Handle system control operations"""
        try:
            if action == 'open_application':
                app_name = params.get('application', params.get('application_name', ''))
                return self._open_application(app_name)
            elif action == 'take_screenshot':
                return self._take_screenshot()
            elif action == 'get_processes':
                return self._get_running_processes()
            elif action == 'get_system_info':
                return self._get_system_info()
            elif action == 'get_services':
                return self._get_windows_services()
            elif action == 'shutdown' or action == 'shutdown_system':
                return self._shutdown_system()
            elif action == 'restart' or action == 'restart_system':
                return self._restart_system()
            elif action == 'lock_screen':
                return self._lock_screen()
            elif action == 'default_action':
                return self._get_system_info()  # Default to system info
            else:
                return f"System control action '{action}' not implemented yet."
        except Exception as e:
            return f"System control failed: {str(e)}"

    def _open_application(self, app_name: str) -> str:
        """Open an application"""
        try:
            if not app_name:
                return "Please specify an application name."
            
            # Check aliases
            app_command = Config.APP_ALIASES.get(app_name.lower(), app_name)
            
            subprocess.Popen([app_command], shell=True)
            return f"Opening {app_name}..."
        except Exception as e:
            return f"Failed to open {app_name}: {str(e)}"

    def _take_screenshot(self) -> str:
        """Take a screenshot"""
        try:
            screenshot = pyautogui.screenshot()
            filename = f"screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            file_path = Config.WORKSPACE_DIR / filename
            screenshot.save(file_path)
            return f"Screenshot saved: {filename}"
        except Exception as e:
            return f"Failed to take screenshot: {str(e)}"

    def _get_running_processes(self) -> str:
        """Get running processes using PowerShell"""
        try:
            script = "Get-Process | Sort-Object CPU -Descending | Select-Object -First 10 | Format-Table Name, Id, CPU -AutoSize"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Top 10 processes by CPU usage:\n{result.stdout}"
        except Exception as e:
            return f"Failed to get processes: {str(e)}"

    def _get_system_info(self) -> str:
        """Get system information"""
        try:
            script = "Get-ComputerInfo | Select-Object WindowsProductName, WindowsVersion, TotalPhysicalMemory, CsProcessors"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"System Information:\n{result.stdout}"
        except Exception as e:
            return f"Failed to get system info: {str(e)}"

    def _get_windows_services(self) -> str:
        """Get Windows services"""
        try:
            script = "Get-Service | Where-Object {$_.Status -eq 'Running'} | Select-Object -First 10 | Format-Table Name, Status -AutoSize"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Running services (first 10):\n{result.stdout}"
        except Exception as e:
            return f"Failed to get services: {str(e)}"

    def _shutdown_system(self) -> str:
        """Shutdown the system"""
        try:
            # Ask for confirmation first
            return "System shutdown requested. For safety, please run 'shutdown /s /t 30' manually to shutdown in 30 seconds, or 'shutdown /a' to cancel."
        except Exception as e:
            return f"Failed to shutdown: {str(e)}"

    def _restart_system(self) -> str:
        """Restart the system"""
        try:
            return "System restart requested. For safety, please run 'shutdown /r /t 30' manually to restart in 30 seconds, or 'shutdown /a' to cancel."
        except Exception as e:
            return f"Failed to restart: {str(e)}"

    def _lock_screen(self) -> str:
        """Lock the screen"""
        try:
            subprocess.run(["rundll32.exe", "user32.dll,LockWorkStation"], check=True)
            return "Screen locked successfully."
        except Exception as e:
            return f"Failed to lock screen: {str(e)}"

    # === WEB BROWSE HANDLERS ===
    def _handle_web_browse(self, action: str, params: Dict) -> str:
        """Handle web browsing operations"""
        try:
            if action == 'web_search':
                query = params.get('search_query', params.get('query', ''))
                return self._perform_web_search(query)
            elif action == 'visit_website':
                url = params.get('url', params.get('website', ''))
                return self._visit_website(url)
            elif action == 'open_website':
                # Handle "open chrome" commands
                app_name = params.get('application', 'chrome')
                if app_name.lower() in ['chrome', 'browser', 'firefox', 'edge']:
                    return self._open_application(app_name)
                else:
                    url = params.get('url', app_name)
                    return self._visit_website(url)
            else:
                return f"Web browsing action '{action}' not implemented yet."
        except Exception as e:
            return f"Web browsing failed: {str(e)}"

    def _perform_web_search(self, query: str) -> str:
        """Perform web search"""
        try:
            if not query:
                return "Please specify what you want to search for."
            
            search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
            webbrowser.open(search_url)
            return f"Searching for: {query}"
        except Exception as e:
            return f"Failed to search: {str(e)}"

    def _visit_website(self, url: str) -> str:
        """Visit a website"""
        try:
            if not url:
                return "Please specify a website URL."
            
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            webbrowser.open(url)
            return f"Opening website: {url}"
        except Exception as e:
            return f"Failed to open website: {str(e)}"

    # === POWERSHELL HANDLERS ===
    def _handle_powershell_task(self, action: str, params: Dict) -> str:
        """Handle PowerShell operations"""
        try:
            if action == 'run_powershell':
                command = params.get('command', '')
                if not command:
                    # Generate PowerShell from natural language
                    instruction = params.get('instruction', 'get system info')
                    command = self.ai_core.generate_powershell_script(instruction)
                
                return self._execute_powershell(command)
            else:
                return f"PowerShell action '{action}' not implemented yet."
        except Exception as e:
            return f"PowerShell operation failed: {str(e)}"

    def _execute_powershell(self, script: str) -> str:
        """Execute PowerShell script"""
        try:
            log_info(f"Executing PowerShell: {script}")
            result = subprocess.run(
                ["powershell", "-Command", script],
                capture_output=True,
                text=True,
                check=True,
                timeout=30
            )
            
            output = result.stdout.strip()
            if result.stderr:
                output += f"\nWarnings: {result.stderr.strip()}"
            
            return output or "Command executed successfully."
        except subprocess.CalledProcessError as e:
            return f"PowerShell script failed: {e.stderr.strip() if e.stderr else 'Unknown error'}"
        except subprocess.TimeoutExpired:
            return "PowerShell script timed out (30 seconds exceeded)."
        except Exception as e:
            return f"Failed to execute PowerShell: {str(e)}"

    # === WEATHER HANDLERS ===
    def _handle_weather_inquiry(self, action: str, params: Dict) -> str:
        """Handle weather inquiries"""
        try:
            city = params.get('city', params.get('city_name', 'London'))
            
            # Simple weather API call (you'd need to implement with a real API)
            return f"I'd love to get weather information for {city}, but I need a weather API key to be configured. You can add services like OpenWeatherMap API to get real weather data."
        except Exception as e:
            return f"Weather inquiry failed: {str(e)}"

    # === KNOWLEDGE HANDLERS ===
    def _handle_knowledge_inquiry(self, action: str, params: Dict) -> str:
        """Handle knowledge inquiries"""
        try:
            topic = params.get('topic', params.get('query', ''))
            
            if not topic:
                return "What would you like to know about?"
            
            # Use AI to provide information
            if self.ai_core.gemini_client:
                try:
                    response = self.ai_core.gemini_client.generate_content(
                        f"Provide a brief, informative explanation about: {topic}"
                    )
                    return response.text.strip()
                except Exception as e:
                    log_error(f"AI knowledge query failed: {e}")
            
            return f"I'd be happy to tell you about {topic}, but I need my AI knowledge system to be properly configured."
        except Exception as e:
            return f"Knowledge inquiry failed: {str(e)}"

    # === MEMORY HANDLERS ===
    def _handle_memory_query(self, action: str, params: Dict) -> str:
        """Handle memory operations"""
        try:
            if action == 'store_memory':
                content = params.get('content', params.get('text', ''))
                return self._store_memory_item(content)
            elif action == 'recall_memory':
                return self._recall_memory()
            elif action == 'recall_facts':
                return self._recall_facts()
            elif action == 'recall_conversation':
                return self._recall_conversation()
            elif action == 'clear_memory':
                return self._clear_memory()
            elif action == 'retrieve_name':
                return self._retrieve_user_name()
            else:
                return f"Memory action '{action}' not implemented yet."
        except Exception as e:
            return f"Memory operation failed: {str(e)}"

    # === CLIPBOARD HANDLERS ===
    def _handle_clipboard_management(self, action: str, params: Dict) -> str:
        """Handle clipboard operations"""
        try:
            if action == 'read_clipboard':
                return self._read_clipboard()
            elif action == 'write_clipboard':
                text = params.get('text', params.get('content', ''))
                return self._write_clipboard(text)
            else:
                return f"Clipboard action '{action}' not implemented yet."
        except Exception as e:
            return f"Clipboard operation failed: {str(e)}"

    def _read_clipboard(self) -> str:
        """Read clipboard content"""
        try:
            content = pyperclip.paste()
            if not content:
                return "Clipboard is empty."
            
            if len(content) > 500:
                content = content[:500] + "... (truncated)"
            
            return f"Clipboard content: {content}"
        except Exception as e:
            return f"Failed to read clipboard: {str(e)}"

    def _write_clipboard(self, text: str) -> str:
        """Write to clipboard"""
        try:
            if not text:
                return "Please specify text to copy to clipboard."
            
            pyperclip.copy(text)
            return f"Copied to clipboard: {text[:100]}..."
        except Exception as e:
            return f"Failed to write to clipboard: {str(e)}"

    # === WINDOWS AUTOMATION HANDLERS ===
    def _handle_windows_automation(self, action: str, params: Dict) -> str:
        """Handle Windows automation"""
        try:
            if action == 'click_coordinates':
                x = params.get('x', 0)
                y = params.get('y', 0)
                return self._click_at_coordinates(x, y)
            elif action == 'send_keys':
                keys = params.get('keys', params.get('text', ''))
                return self._send_keys(keys)
            elif action == 'get_window_list':
                return self._get_window_list()
            else:
                return f"Windows automation action '{action}' not implemented yet."
        except Exception as e:
            return f"Windows automation failed: {str(e)}"

    def _click_at_coordinates(self, x: int, y: int) -> str:
        """Click at specified coordinates"""
        try:
            pyautogui.click(x, y)
            return f"Clicked at coordinates ({x}, {y})"
        except Exception as e:
            return f"Failed to click: {str(e)}"

    def _send_keys(self, keys: str) -> str:
        """Send keystrokes"""
        try:
            if not keys:
                return "Please specify keys to send."
            
            pyautogui.typewrite(keys)
            return f"Sent keys: {keys}"
        except Exception as e:
            return f"Failed to send keys: {str(e)}"

    def _get_window_list(self) -> str:
        """Get list of open windows"""
        try:
            # This is a simplified version - you might want to use pygetwindow for more functionality
            script = "Get-Process | Where-Object {$_.MainWindowTitle -ne ''} | Select-Object -First 10 ProcessName, MainWindowTitle | Format-Table -AutoSize"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Open windows:\n{result.stdout}"
        except Exception as e:
            return f"Failed to get window list: {str(e)}"

    # === SYSTEM STATUS HANDLERS ===
    def _handle_system_status(self, action: str, params: Dict) -> str:
        """Handle system status queries"""
        try:
            if action == 'ram_usage' or action == 'get_memory_usage':
                return self._get_memory_usage()
            elif action == 'cpu_usage' or action == 'get_cpu_usage':
                return self._get_cpu_usage()
            elif action == 'disk_usage' or action == 'get_disk_usage':
                return self._get_disk_usage()
            elif action == 'battery_status' or action == 'get_battery_status':
                return self._get_battery_status()
            elif action == 'network_status' or action == 'get_network_status':
                return self._get_network_status()
            else:
                return self._get_system_status_overview()
        except Exception as e:
            return f"System status query failed: {str(e)}"

    def _get_memory_usage(self) -> str:
        """Get RAM usage information"""
        try:
            script = """
            $mem = Get-WmiObject -Class Win32_OperatingSystem
            $total = [math]::Round($mem.TotalVisibleMemorySize/1KB, 2)
            $free = [math]::Round($mem.FreePhysicalMemory/1KB, 2)
            $used = [math]::Round($total - $free, 2)
            $percent = [math]::Round(($used/$total)*100, 1)
            Write-Output "RAM Usage: $used GB / $total GB ($percent% used)"
            Write-Output "Available: $free GB"
            """
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return result.stdout.strip()
        except Exception as e:
            return f"Failed to get memory usage: {str(e)}"

    def _get_cpu_usage(self) -> str:
        """Get CPU usage information"""
        try:
            script = """
            $cpu = Get-WmiObject -Class Win32_Processor
            $usage = (Get-Counter "\\Processor(_Total)\\% Processor Time").CounterSamples.CookedValue
            $usage = [math]::Round(100 - $usage, 1)
            Write-Output "CPU: $($cpu.Name)"
            Write-Output "Current Usage: $usage%"
            Write-Output "Cores: $($cpu.NumberOfCores) | Logical Processors: $($cpu.NumberOfLogicalProcessors)"
            """
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return result.stdout.strip()
        except Exception as e:
            return f"Failed to get CPU usage: {str(e)}"

    def _get_disk_usage(self) -> str:
        """Get disk usage information"""
        try:
            script = """
            Get-WmiObject -Class Win32_LogicalDisk | Where-Object {$_.DriveType -eq 3} | ForEach-Object {
                $total = [math]::Round($_.Size/1GB, 1)
                $free = [math]::Round($_.FreeSpace/1GB, 1)
                $used = [math]::Round($total - $free, 1)
                $percent = [math]::Round(($used/$total)*100, 1)
                Write-Output "Drive $($_.DeviceID) $used GB / $total GB ($percent% used) | Free: $free GB"
            }
            """
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Disk Usage:\n{result.stdout}"
        except Exception as e:
            return f"Failed to get disk usage: {str(e)}"

    def _get_battery_status(self) -> str:
        """Get battery status"""
        try:
            script = """
            $battery = Get-WmiObject -Class Win32_Battery
            if ($battery) {
                $percent = $battery.EstimatedChargeRemaining
                $status = switch ($battery.BatteryStatus) {
                    1 { "Discharging" }
                    2 { "Charging" }
                    3 { "Critical" }
                    4 { "Low" }
                    5 { "High" }
                    6 { "Charging Complete" }
                    default { "Unknown" }
                }
                Write-Output "Battery: $percent% - $status"
            } else {
                Write-Output "No battery detected (Desktop PC)"
            }
            """
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return result.stdout.strip()
        except Exception as e:
            return f"Failed to get battery status: {str(e)}"

    def _get_network_status(self) -> str:
        """Get network status"""
        try:
            script = """
            Get-NetAdapter | Where-Object {$_.Status -eq 'Up'} | ForEach-Object {
                $speed = if ($_.LinkSpeed) { 
                    [math]::Round($_.LinkSpeed/1MB, 0).ToString() + " Mbps" 
                } else { "Unknown" }
                Write-Output "$($_.Name): $($_.InterfaceDescription) - $speed"
            }
            """
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Active Network Adapters:\n{result.stdout}"
        except Exception as e:
            return f"Failed to get network status: {str(e)}"

    def _get_system_status_overview(self) -> str:
        """Get comprehensive system status"""
        try:
            memory = self._get_memory_usage()
            cpu = self._get_cpu_usage()
            disk = self._get_disk_usage()
            return f"System Status Overview:\n\n{cpu}\n\n{memory}\n\n{disk}"
        except Exception as e:
            return f"Failed to get system overview: {str(e)}"

    # === NETWORK OPERATIONS HANDLERS ===
    def _handle_network_operations(self, action: str, params: Dict) -> str:
        """Handle network operations"""
        try:
            if action == 'show_network_status' or action == 'get_network_status':
                return self._get_network_status()
            elif action == 'disconnect_wifi':
                return self._disconnect_wifi()
            elif action == 'connect_wifi':
                return self._connect_wifi()
            elif action == 'ping_host':
                host = params.get('host', 'google.com')
                return self._ping_host(host)
            elif action == 'show_ip_config':
                return self._show_ip_config()
            else:
                return f"Network operation '{action}' not implemented yet."
        except Exception as e:
            return f"Network operation failed: {str(e)}"

    def _disconnect_wifi(self) -> str:
        """Disconnect from WiFi"""
        try:
            script = "netsh wlan disconnect"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return "WiFi disconnected successfully."
        except Exception as e:
            return f"Failed to disconnect WiFi: {str(e)}"

    def _connect_wifi(self) -> str:
        """Connect to WiFi (shows available networks)"""
        try:
            script = "netsh wlan show profiles"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"Available WiFi profiles:\n{result.stdout}"
        except Exception as e:
            return f"Failed to show WiFi profiles: {str(e)}"

    def _ping_host(self, host: str) -> str:
        """Ping a host"""
        try:
            result = subprocess.run(["ping", "-n", "4", host], 
                                  capture_output=True, text=True, check=True)
            return f"Ping results for {host}:\n{result.stdout}"
        except Exception as e:
            return f"Failed to ping {host}: {str(e)}"

    def _show_ip_config(self) -> str:
        """Show IP configuration"""
        try:
            script = "ipconfig /all"
            result = subprocess.run(["powershell", "-Command", script], 
                                  capture_output=True, text=True, check=True)
            return f"IP Configuration:\n{result.stdout[:2000]}..."  # Limit output
        except Exception as e:
            return f"Failed to get IP configuration: {str(e)}"

    # === MEDIA CONTROL HANDLERS ===
    def _handle_media_control(self, action: str, params: Dict) -> str:
        """Handle media control operations"""
        try:
            if action == 'turn_on_music' or action == 'play_music':
                return self._play_music()
            elif action == 'pause_music':
                return self._pause_music()
            elif action == 'stop_music':
                return self._stop_music()
            elif action == 'next_track':
                return self._next_track()
            elif action == 'previous_track':
                return self._previous_track()
            elif action == 'set_volume':
                volume = params.get('volume', 50)
                return self._set_volume(volume)
            else:
                return f"Media control action '{action}' not implemented yet."
        except Exception as e:
            return f"Media control failed: {str(e)}"

    def _play_music(self) -> str:
        """Play music (opens default music app)"""
        try:
            # Try to open Spotify first, then Windows Media Player
            music_apps = ['spotify', 'wmplayer', 'groove']
            
            for app in music_apps:
                try:
                    subprocess.Popen([app], shell=True)
                    return f"Opening {app} for music playback..."
                except:
                    continue
            
            return "Could not find a music application to open. Please install Spotify or Windows Media Player."
        except Exception as e:
            return f"Failed to start music: {str(e)}"

    def _pause_music(self) -> str:
        """Pause music using media keys"""
        try:
            pyautogui.press('playpause')
            return "Music paused/resumed."
        except Exception as e:
            return f"Failed to pause music: {str(e)}"

    def _stop_music(self) -> str:
        """Stop music"""
        try:
            pyautogui.press('stop')
            return "Music stopped."
        except Exception as e:
            return f"Failed to stop music: {str(e)}"

    def _next_track(self) -> str:
        """Next track"""
        try:
            pyautogui.press('nexttrack')
            return "Skipped to next track."
        except Exception as e:
            return f"Failed to skip track: {str(e)}"

    def _previous_track(self) -> str:
        """Previous track"""
        try:
            pyautogui.press('prevtrack')
            return "Went to previous track."
        except Exception as e:
            return f"Failed to go to previous track: {str(e)}"

    def _set_volume(self, volume: int) -> str:
        """Set system volume"""
        try:
            # Volume keys approach (simplified)
            if volume > 50:
                for _ in range(5):
                    pyautogui.press('volumeup')
            else:
                for _ in range(5):
                    pyautogui.press('volumedown')
            
            return f"Volume adjusted to approximately {volume}%"
        except Exception as e:
            return f"Failed to set volume: {str(e)}"

    # === MEMORY SYSTEM ===
    def _load_memory(self) -> Dict:
        """Load memory from file"""
        try:
            if os.path.exists(self.memory_file):
                with open(self.memory_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            log_error(f"Failed to load memory: {e}")
        
        return {
            "facts": [],
            "conversations": [],
            "preferences": {},
            "interactions": []
        }

    def _save_memory(self):
        """Save memory to file"""
        try:
            with open(self.memory_file, 'w', encoding='utf-8') as f:
                json.dump(self.memory, f, indent=2, ensure_ascii=False)
        except Exception as e:
            log_error(f"Failed to save memory: {e}")

    def _store_memory_item(self, content: str) -> str:
        """Store a memory item"""
        try:
            memory_item = {
                "content": content,
                "timestamp": datetime.now().isoformat(),
                "session_id": self.session_id
            }
            
            if "facts" not in self.memory:
                self.memory["facts"] = []
            
            self.memory["facts"].append(memory_item)
            self.memory["facts"] = self.memory["facts"][-50:]  # Keep last 50 facts
            
            self._save_memory()
            return f"Stored in memory: {content}"
        except Exception as e:
            return f"Failed to store memory: {str(e)}"

    def _recall_memory(self) -> str:
        """Recall recent memory items"""
        try:
            if "facts" not in self.memory or not self.memory["facts"]:
                return "I don't have any stored memories yet."
            
            recent_facts = self.memory["facts"][-5:]  # Last 5 facts
            memory_text = "Here's what I remember:\n"
            
            for i, fact in enumerate(recent_facts, 1):
                memory_text += f"{i}. {fact['content']}\n"
            
            return memory_text
        except Exception as e:
            return f"Failed to recall memory: {str(e)}"

    def _recall_facts(self) -> str:
        """Recall stored facts"""
        return self._recall_memory()  # Same as recall_memory for now

    def _recall_conversation(self) -> str:
        """Recall recent conversation"""
        try:
            if not self.conversation_history:
                return "No conversation history in this session."
            
            history_text = "Recent conversation:\n"
            for entry in self.conversation_history[-6:]:  # Last 6 exchanges
                role = "You" if entry["role"] == "user" else "Me"
                content = entry["content"][:100] + "..." if len(entry["content"]) > 100 else entry["content"]
                history_text += f"{role}: {content}\n"
            
            return history_text
        except Exception as e:
            return f"Failed to recall conversation: {str(e)}"

    def _clear_memory(self) -> str:
        """Clear stored memory"""
        try:
            self.memory = {
                "facts": [],
                "conversations": [],
                "preferences": {},
                "interactions": []
            }
            self._save_memory()
            return "Memory cleared successfully."
        except Exception as e:
            return f"Failed to clear memory: {str(e)}"

    def _store_interaction(self, command: str, response: str, analysis: CommandAnalysis):
        """Store interaction in memory"""
        try:
            interaction = {
                "command": command,
                "response": response,
                "intent": analysis.intent,
                "action": analysis.action,
                "timestamp": datetime.now().isoformat(),
                "session_id": self.session_id
            }
            
            if "interactions" not in self.memory:
                self.memory["interactions"] = []
            
            self.memory["interactions"].append(interaction)
            self.memory["interactions"] = self.memory["interactions"][-100:]  # Keep last 100
            
            self._save_memory()
        except Exception as e:
            log_error(f"Failed to store interaction: {e}")

    def _update_conversation_history(self, command: str, response: str):
        """Update conversation history"""
        self.conversation_history.append({"role": "user", "content": command})
        self.conversation_history.append({"role": "assistant", "content": response})
        self.conversation_history = self.conversation_history[-10:]  # Keep last 10 entries

    def _get_comprehensive_help(self) -> str:
        """Get comprehensive help text"""
        return """
        JARVIS Enhanced Commands:
        
        📁 File Operations:
        - "Create a Word document about AI"
        - "List files in my workspace"
        - "Read file config.txt"
        - "Delete file old_document.txt"
        - "Take a screenshot"
        
        💻 System Control:
        - "Open Calculator"
        - "Open Chrome browser"
        - "Show running processes"
        - "Get system information"
        
        🌐 Web & Search:
        - "Search for artificial intelligence"
        - "Open Google website"
        - "Visit github.com"
        
        🧠 Memory & Knowledge:
        - "Remember that I prefer dark mode"
        - "What do you remember about me?"
        - "Tell me about quantum computing"
        - "What did we discuss earlier?"
        
        📋 Clipboard:
        - "Read my clipboard"
        - "Copy 'Hello World' to clipboard"
        
        🖱️ Automation:
        - "Click at coordinates 500, 300"
        - "Type 'Hello World'"
        - "Show open windows"
        
        ⚡ PowerShell:
        - "Run PowerShell to get system info"
        - "Execute command to list processes"
        
        Just speak naturally - I'll understand what you want to do!
        """

    def cleanup(self):
        """Clean up resources"""
        try:
            self._save_memory()
            log_info("JARVIS Core cleanup completed.")
        except Exception as e:
            log_error(f"Cleanup failed: {e}")